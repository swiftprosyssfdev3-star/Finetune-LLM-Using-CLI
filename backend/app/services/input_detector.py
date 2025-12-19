"""
Flexible Input Detection Service
Automatically detect and parse any input format

Supports:
- Image folders/archives (jpg, png, tiff, pdf, webp, bmp)
- Data files (xlsx, xls, csv, tsv, json, jsonl, parquet)
- Document files (docx, doc, md, txt) for ground truth
- Schema files (json with schema definition)
- Archives (zip, tar.gz with mixed content)
- Entire folder uploads with recursive scanning
"""

from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple, Union
from dataclasses import dataclass, asdict, field
import json
import re
import zipfile
import tarfile
import tempfile
import shutil
import os


@dataclass
class ImageSet:
    """Detected image set."""
    paths: List[str]
    count: int
    formats: List[str]
    sample_dimensions: Optional[Tuple[int, int]]
    total_size_mb: float
    folder_structure: Optional[Dict[str, int]] = None  # folder -> count mapping


@dataclass
class DataFile:
    """Detected data file."""
    path: str
    format: str  # xlsx, csv, json, jsonl, docx, md
    row_count: int
    column_count: int
    columns: List[str]
    sample_rows: List[Dict[str, Any]]
    detected_schema: Dict[str, str]
    raw_content: Optional[str] = None  # For docx/md files


@dataclass
class OutputSchema:
    """Output schema definition."""
    schema: Dict[str, Any]
    source: str  # auto, file, manual
    sample_output: str


@dataclass
class ProcessingInstructions:
    """Instructions for CLI agent processing."""
    task_type: str  # ocr, classification, extraction, qa
    input_format: str
    output_format: str
    matching_strategy: str  # filename, index, metadata
    special_instructions: List[str]


@dataclass
class ProjectInput:
    """Complete project input."""
    images: Optional[ImageSet]
    data: Optional[DataFile]
    schema: Optional[OutputSchema]
    processing: Optional[ProcessingInstructions]
    warnings: List[str]
    suggestions: List[str]
    agent_prompt: Optional[str] = None  # Auto-generated prompt for CLI agent


class InputDetector:
    """
    Automatically detect and parse any input format.

    Supports:
    - Image folders/archives (jpg, png, tiff, pdf, webp, bmp)
    - Data files (xlsx, xls, csv, tsv, json, jsonl, parquet)
    - Document files (docx, doc, md, txt) for ground truth
    - Schema files (json with schema definition)
    - Archives (zip, tar.gz with mixed content)
    - Entire folder uploads with recursive scanning
    """

    IMAGE_EXTENSIONS = {
        '.jpg', '.jpeg', '.png', '.tiff', '.tif',
        '.bmp', '.webp', '.pdf', '.gif', '.heic', '.heif'
    }

    DATA_EXTENSIONS = {
        '.xlsx', '.xls', '.csv', '.tsv',
        '.json', '.jsonl', '.parquet'
    }

    DOCUMENT_EXTENSIONS = {
        '.docx', '.doc', '.md', '.markdown', '.txt', '.rtf'
    }

    ARCHIVE_EXTENSIONS = {
        '.zip', '.tar', '.gz', '.tgz', '.tar.gz', '.7z', '.rar'
    }

    def __init__(self, upload_dir: str = "./uploads"):
        """
        Initialize the input detector.

        Args:
            upload_dir: Directory for uploaded/extracted files
        """
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def detect_inputs(self, paths: List[str]) -> ProjectInput:
        """
        Detect and parse all inputs from files or folders.

        Args:
            paths: List of file/folder paths

        Returns:
            ProjectInput with detected images, data, schema, and processing instructions
        """

        images = None
        data = None
        schema = None
        warnings = []
        suggestions = []

        all_images = []
        all_data_files = []
        folder_structure = {}

        for path_str in paths:
            path = Path(path_str)

            if not path.exists():
                warnings.append(f"Path not found: {path}")
                continue

            if path.is_dir():
                # Recursively scan entire folder
                result = await self._scan_folder_recursive(path)
                if result['images']:
                    all_images.extend(result['images'])
                    # Track folder structure
                    for img_path in result['images']:
                        folder = str(Path(img_path).parent.relative_to(path) if path in Path(img_path).parents else Path(img_path).parent.name)
                        folder_structure[folder] = folder_structure.get(folder, 0) + 1
                if result['data_files']:
                    all_data_files.extend(result['data_files'])
                if result['schema'] and schema is None:
                    schema = result['schema']

            elif self._is_archive(path):
                # Extract and scan archive
                try:
                    extracted = await self._extract_archive(path)
                    result = await self._scan_folder_recursive(extracted)
                    if result['images']:
                        all_images.extend(result['images'])
                    if result['data_files']:
                        all_data_files.extend(result['data_files'])
                except Exception as e:
                    warnings.append(f"Failed to extract archive {path}: {str(e)}")

            elif path.suffix.lower() in self.IMAGE_EXTENSIONS:
                all_images.append(str(path))

            elif path.suffix.lower() in self.DATA_EXTENSIONS:
                all_data_files.append(path)

            elif path.suffix.lower() in self.DOCUMENT_EXTENSIONS:
                all_data_files.append(path)

            elif path.suffix.lower() == '.json':
                try:
                    content = json.loads(path.read_text(encoding='utf-8'))
                    if self._is_schema_file(path, content):
                        schema = OutputSchema(
                            schema=content,
                            source='file',
                            sample_output=json.dumps(content, indent=2)[:500]
                        )
                    else:
                        all_data_files.append(path)
                except Exception as e:
                    warnings.append(f"Failed to parse JSON file {path}: {str(e)}")

        # Consolidate images
        if all_images:
            formats = list(set(Path(p).suffix.lower() for p in all_images))
            total_size = 0
            for p in all_images:
                try:
                    total_size += Path(p).stat().st_size
                except:
                    pass

            sample_dims = self._get_image_dimensions(Path(all_images[0])) if all_images else None

            images = ImageSet(
                paths=sorted(all_images),
                count=len(all_images),
                formats=formats,
                sample_dimensions=sample_dims,
                total_size_mb=round(total_size / (1024 * 1024), 2),
                folder_structure=folder_structure if folder_structure else None
            )

        # Parse best data file
        if all_data_files:
            data = await self._select_and_parse_data(all_data_files)

        # Auto-generate schema from data
        if schema is None and data is not None:
            schema = self._generate_schema(data)
            suggestions.append(
                "Schema auto-generated from data columns. "
                "Review and adjust field types if needed."
            )

        # Generate processing instructions and suggestions
        processing = self._generate_processing_instructions(images, data, schema)

        # Validate matching and add suggestions
        if images and data:
            self._validate_matching(images, data, suggestions, warnings)

        # Generate agent prompt
        agent_prompt = self._generate_agent_prompt(images, data, schema, processing)

        return ProjectInput(
            images=images,
            data=data,
            schema=schema,
            processing=processing,
            warnings=warnings,
            suggestions=suggestions,
            agent_prompt=agent_prompt
        )

    async def _scan_folder_recursive(self, folder: Path) -> Dict[str, Any]:
        """Recursively scan a folder for all supported files."""

        result = {
            'images': [],
            'data_files': [],
            'schema': None
        }

        # Walk through all files recursively
        for root, dirs, files in os.walk(folder):
            # Skip hidden directories
            dirs[:] = [d for d in dirs if not d.startswith('.')]

            for file in files:
                # Skip hidden files
                if file.startswith('.'):
                    continue

                file_path = Path(root) / file
                ext = file_path.suffix.lower()

                if ext in self.IMAGE_EXTENSIONS:
                    result['images'].append(str(file_path))

                elif ext in self.DATA_EXTENSIONS or ext in self.DOCUMENT_EXTENSIONS:
                    result['data_files'].append(file_path)

                elif ext == '.json':
                    try:
                        content = json.loads(file_path.read_text(encoding='utf-8'))
                        if self._is_schema_file(file_path, content) and result['schema'] is None:
                            result['schema'] = OutputSchema(
                                schema=content,
                                source='file',
                                sample_output=json.dumps(content, indent=2)[:500]
                            )
                        else:
                            result['data_files'].append(file_path)
                    except:
                        pass

        return result

    async def _select_and_parse_data(self, data_files: List[Path]) -> Optional[DataFile]:
        """Select the best data file and parse it."""

        if not data_files:
            return None

        # Priority: xlsx > csv > json > docx > md
        def priority(f: Path) -> int:
            ext = f.suffix.lower()
            if ext in ['.xlsx', '.xls']:
                return 0
            elif ext == '.csv':
                return 1
            elif ext in ['.json', '.jsonl']:
                return 2
            elif ext in ['.docx', '.doc']:
                return 3
            elif ext in ['.md', '.markdown', '.txt']:
                return 4
            else:
                return 5

        data_files.sort(key=priority)

        for data_file in data_files:
            try:
                return await self._parse_data_file(data_file)
            except Exception as e:
                continue

        return None

    async def _parse_data_file(self, path: Path) -> DataFile:
        """Parse a data file and extract schema information."""

        ext = path.suffix.lower()
        df = None
        rows = []
        columns = []
        raw_content = None

        # Handle document files (DOCX, MD, TXT)
        if ext in self.DOCUMENT_EXTENSIONS:
            return await self._parse_document_file(path)

        try:
            import pandas as pd

            if ext in ('.xlsx', '.xls'):
                df = pd.read_excel(path)
            elif ext == '.csv':
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        df = pd.read_csv(path, encoding=encoding)
                        break
                    except:
                        continue
            elif ext == '.tsv':
                df = pd.read_csv(path, sep='\t')
            elif ext == '.jsonl':
                df = pd.read_json(path, lines=True)
            elif ext == '.json':
                content = json.loads(path.read_text(encoding='utf-8'))
                if isinstance(content, list):
                    df = pd.DataFrame(content)
                elif isinstance(content, dict) and 'data' in content:
                    df = pd.DataFrame(content['data'])
                elif isinstance(content, dict) and 'records' in content:
                    df = pd.DataFrame(content['records'])
                elif isinstance(content, dict) and 'items' in content:
                    df = pd.DataFrame(content['items'])
                else:
                    df = pd.DataFrame([content])
            elif ext == '.parquet':
                df = pd.read_parquet(path)
            else:
                raise ValueError(f"Unsupported format: {ext}")

            if df is not None:
                columns = list(df.columns)
                rows = df.head(10).to_dict('records')

                # Clean NaN values
                for row in rows:
                    for k, v in row.items():
                        if pd.isna(v):
                            row[k] = None

        except ImportError:
            # Fallback without pandas
            rows, columns = self._parse_without_pandas(path, ext)

        # Detect column types
        detected_schema = {}
        for col in columns:
            sample_values = [r.get(col) for r in rows if r.get(col) is not None]
            detected_schema[col] = self._detect_type(sample_values)

        row_count = len(df) if df is not None else len(rows)

        return DataFile(
            path=str(path),
            format=ext.replace('.', ''),
            row_count=row_count,
            column_count=len(columns),
            columns=columns,
            sample_rows=rows,
            detected_schema=detected_schema,
            raw_content=raw_content
        )

    async def _parse_document_file(self, path: Path) -> DataFile:
        """Parse document files (DOCX, MD, TXT) as ground truth."""

        ext = path.suffix.lower()
        content = ""
        rows = []
        columns = ['content', 'source_file']

        if ext in ['.docx', '.doc']:
            content = self._parse_docx(path)
        elif ext in ['.md', '.markdown']:
            content = self._parse_markdown(path)
        elif ext == '.txt':
            content = path.read_text(encoding='utf-8', errors='ignore')
        elif ext == '.rtf':
            content = self._parse_rtf(path)

        # Try to extract structured data from content
        extracted_rows = self._extract_structured_data(content, path.name)

        if extracted_rows:
            rows = extracted_rows
            if rows:
                columns = list(rows[0].keys())
        else:
            # Use raw content as single record
            rows = [{'content': content, 'source_file': path.name}]

        # Detect schema
        detected_schema = {}
        for col in columns:
            sample_values = [r.get(col) for r in rows[:10] if r.get(col) is not None]
            detected_schema[col] = self._detect_type(sample_values)

        return DataFile(
            path=str(path),
            format=ext.replace('.', ''),
            row_count=len(rows),
            column_count=len(columns),
            columns=columns,
            sample_rows=rows[:10],
            detected_schema=detected_schema,
            raw_content=content[:5000] if len(content) > 5000 else content
        )

    def _parse_docx(self, path: Path) -> str:
        """Parse DOCX file content."""
        try:
            from docx import Document
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    tables_text.append(row_text)

            content = '\n'.join(paragraphs)
            if tables_text:
                content += '\n\n--- Tables ---\n' + '\n'.join(tables_text)

            return content
        except ImportError:
            # Fallback: try to extract raw text
            try:
                import zipfile
                with zipfile.ZipFile(path, 'r') as zf:
                    xml_content = zf.read('word/document.xml').decode('utf-8')
                    # Simple regex to extract text
                    text = re.sub(r'<[^>]+>', '', xml_content)
                    return text
            except:
                return f"[DOCX file: {path.name} - install python-docx to parse]"
        except Exception as e:
            return f"[Error parsing DOCX: {str(e)}]"

    def _parse_markdown(self, path: Path) -> str:
        """Parse Markdown file content."""
        try:
            content = path.read_text(encoding='utf-8', errors='ignore')
            return content
        except Exception as e:
            return f"[Error parsing Markdown: {str(e)}]"

    def _parse_rtf(self, path: Path) -> str:
        """Parse RTF file content."""
        try:
            from striprtf.striprtf import rtf_to_text
            content = path.read_text(encoding='utf-8', errors='ignore')
            return rtf_to_text(content)
        except ImportError:
            # Basic RTF parsing
            content = path.read_text(encoding='utf-8', errors='ignore')
            # Remove RTF control words
            text = re.sub(r'\\[a-z]+\d* ?', '', content)
            text = re.sub(r'[{}]', '', text)
            return text
        except Exception as e:
            return f"[Error parsing RTF: {str(e)}]"

    def _extract_structured_data(self, content: str, filename: str) -> List[Dict[str, Any]]:
        """Try to extract structured data from document content."""

        rows = []

        # Try to detect table-like structure (| separated or tab separated)
        lines = content.strip().split('\n')

        # Check for markdown table
        if any('|' in line for line in lines[:10]):
            rows = self._parse_markdown_table(lines)
            if rows:
                return rows

        # Check for tab-separated values
        if any('\t' in line for line in lines[:10]):
            rows = self._parse_tsv_content(lines)
            if rows:
                return rows

        # Check for key-value pairs
        kv_pattern = r'^([^:]+):\s*(.+)$'
        kv_rows = []
        current_record = {}

        for line in lines:
            match = re.match(kv_pattern, line.strip())
            if match:
                key, value = match.groups()
                key = key.strip().lower().replace(' ', '_')
                current_record[key] = value.strip()
            elif line.strip() == '' and current_record:
                kv_rows.append(current_record)
                current_record = {}

        if current_record:
            kv_rows.append(current_record)

        if kv_rows and len(kv_rows[0]) > 1:
            return kv_rows

        # Check for numbered list items
        numbered_pattern = r'^\d+[\.\)]\s*(.+)$'
        numbered_items = []
        for line in lines:
            match = re.match(numbered_pattern, line.strip())
            if match:
                numbered_items.append({'item': match.group(1), 'source': filename})

        if len(numbered_items) > 3:
            return numbered_items

        return []

    def _parse_markdown_table(self, lines: List[str]) -> List[Dict[str, Any]]:
        """Parse markdown table from lines."""
        rows = []
        headers = []
        in_table = False

        for line in lines:
            if '|' in line:
                cells = [c.strip() for c in line.split('|')]
                cells = [c for c in cells if c]  # Remove empty

                if not headers:
                    headers = cells
                    in_table = True
                elif re.match(r'^[\-:| ]+$', line):
                    # Skip separator line
                    continue
                elif in_table and cells:
                    row = dict(zip(headers, cells))
                    rows.append(row)
            elif in_table and not line.strip():
                # End of table
                break

        return rows

    def _parse_tsv_content(self, lines: List[str]) -> List[Dict[str, Any]]:
        """Parse tab-separated content."""
        if not lines:
            return []

        headers = lines[0].split('\t')
        headers = [h.strip() for h in headers]

        rows = []
        for line in lines[1:]:
            if line.strip():
                values = line.split('\t')
                row = dict(zip(headers, [v.strip() for v in values]))
                rows.append(row)

        return rows

    def _parse_without_pandas(self, path: Path, ext: str) -> Tuple[List[Dict], List[str]]:
        """Parse data files without pandas."""
        rows = []
        columns = []

        if ext == '.json':
            content = json.loads(path.read_text(encoding='utf-8'))
            if isinstance(content, list) and content:
                columns = list(content[0].keys())
                rows = content[:10]
            elif isinstance(content, dict):
                columns = list(content.keys())
                rows = [content]

        elif ext == '.jsonl':
            with open(path, encoding='utf-8') as f:
                for i, line in enumerate(f):
                    if i >= 10:
                        break
                    row = json.loads(line)
                    rows.append(row)
                    if not columns:
                        columns = list(row.keys())

        elif ext == '.csv':
            with open(path, encoding='utf-8') as f:
                lines = f.readlines()[:11]
                if lines:
                    columns = lines[0].strip().split(',')
                    for line in lines[1:]:
                        values = line.strip().split(',')
                        rows.append(dict(zip(columns, values)))

        return rows, columns

    def _detect_type(self, values: List[Any]) -> str:
        """Detect the type of a column based on sample values."""

        if not values:
            return 'string'

        # Check for dates
        date_patterns = [
            r'^\d{4}-\d{2}-\d{2}$',
            r'^\d{2}/\d{2}/\d{4}$',
            r'^\d{2}\.\d{2}\.\d{4}$',
            r'^\d{4}/\d{2}/\d{2}$',
            r'^\d{1,2}\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}$',
        ]

        str_values = [str(v) for v in values if v is not None]

        for pattern in date_patterns:
            if str_values and all(re.match(pattern, v, re.IGNORECASE) for v in str_values if v):
                return 'date'

        # Check for numbers
        int_count = 0
        float_count = 0
        for v in values:
            if isinstance(v, bool):
                continue
            if isinstance(v, int):
                int_count += 1
            elif isinstance(v, float):
                float_count += 1
            elif isinstance(v, str):
                try:
                    int(v)
                    int_count += 1
                except ValueError:
                    try:
                        float(v)
                        float_count += 1
                    except ValueError:
                        pass

        if int_count == len(values):
            return 'integer'
        if int_count + float_count == len(values):
            return 'number'

        return 'string'

    def _generate_schema(self, data: DataFile) -> OutputSchema:
        """Auto-generate JSON schema from data."""

        schema = {
            "type": "object",
            "properties": {
                "records": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {}
                    }
                }
            }
        }

        type_mapping = {
            'integer': 'integer',
            'number': 'number',
            'date': 'string',
            'string': 'string'
        }

        for col, dtype in data.detected_schema.items():
            json_type = type_mapping.get(dtype, 'string')

            prop = {"type": json_type}
            if dtype == 'date':
                prop["format"] = "date"
                prop["description"] = f"{col} (YYYY-MM-DD format)"

            schema["properties"]["records"]["items"]["properties"][col] = prop

        # Create sample output
        sample = {"records": data.sample_rows[:1]} if data.sample_rows else {"records": []}

        return OutputSchema(
            schema=schema,
            source='auto',
            sample_output=json.dumps(sample, indent=2, default=str)
        )

    def _generate_processing_instructions(
        self,
        images: Optional[ImageSet],
        data: Optional[DataFile],
        schema: Optional[OutputSchema]
    ) -> Optional[ProcessingInstructions]:
        """Generate processing instructions for the CLI agent."""

        if not images:
            return None

        # Determine task type
        task_type = "ocr"  # Default
        if data and data.columns:
            cols_lower = [c.lower() for c in data.columns]
            if any('class' in c or 'label' in c or 'category' in c for c in cols_lower):
                task_type = "classification"
            elif any('question' in c or 'answer' in c or 'qa' in c for c in cols_lower):
                task_type = "qa"
            elif len(data.columns) > 3:
                task_type = "extraction"

        # Determine matching strategy
        matching_strategy = "filename"
        if data and images:
            # Check if filenames match
            image_names = set(Path(p).stem.lower() for p in images.paths)
            if data.columns and 'filename' in [c.lower() for c in data.columns]:
                matching_strategy = "filename_column"
            elif data.columns and any('image' in c.lower() for c in data.columns):
                matching_strategy = "image_column"
            elif images.count == data.row_count:
                matching_strategy = "index"

        # Special instructions
        special_instructions = []

        if images.folder_structure:
            folders = list(images.folder_structure.keys())
            if len(folders) > 1:
                special_instructions.append(
                    f"Images organized in {len(folders)} folders: {', '.join(folders[:5])}"
                )

        if data and data.raw_content:
            special_instructions.append(
                "Ground truth provided as document text - parse and match to images"
            )

        if data and data.format in ['docx', 'md', 'txt']:
            special_instructions.append(
                f"Ground truth in {data.format.upper()} format - extract structured data"
            )

        return ProcessingInstructions(
            task_type=task_type,
            input_format=f"images ({', '.join(images.formats)})",
            output_format="json" if schema else "auto-detect",
            matching_strategy=matching_strategy,
            special_instructions=special_instructions
        )

    def _validate_matching(
        self,
        images: ImageSet,
        data: DataFile,
        suggestions: List[str],
        warnings: List[str]
    ):
        """Validate image-data matching and add suggestions."""

        if images.count == data.row_count:
            suggestions.append(
                f"Perfect match: {images.count} images and {data.row_count} data rows."
            )
        elif images.count != data.row_count:
            ratio = data.row_count / images.count if images.count > 0 else 0

            if ratio > 1 and abs(ratio - round(ratio)) < 0.1:
                suggestions.append(
                    f"Detected ~{int(round(ratio))} records per image. "
                    f"Multi-record extraction mode will be used."
                )
            elif 0.8 < ratio < 1.2:
                suggestions.append(
                    f"Approximate match ({images.count} images, {data.row_count} rows). "
                    f"Agent will match by filename or content."
                )
            else:
                warnings.append(
                    f"Significant mismatch: {images.count} images vs {data.row_count} data rows. "
                    f"Please verify your data or the agent will attempt automatic matching."
                )

    def _generate_agent_prompt(
        self,
        images: Optional[ImageSet],
        data: Optional[DataFile],
        schema: Optional[OutputSchema],
        processing: Optional[ProcessingInstructions]
    ) -> Optional[str]:
        """Generate an auto-prompt for the CLI agent."""

        if not images:
            return None

        prompt_parts = [
            "# VLM Fine-Tuning Task",
            "",
            "## Input Data Analysis",
            f"- **Images**: {images.count} files ({', '.join(images.formats)})",
        ]

        if images.folder_structure:
            prompt_parts.append(f"- **Folder structure**: {json.dumps(images.folder_structure)}")

        if data:
            prompt_parts.extend([
                f"- **Ground truth**: {data.format.upper()} with {data.row_count} records",
                f"- **Columns**: {', '.join(data.columns[:10])}",
            ])

        prompt_parts.extend([
            "",
            "## Task Instructions",
            "",
            "1. **Analyze** the input images and ground truth data",
            "2. **Match** images to their corresponding ground truth records",
            "3. **Create** a data loader that properly pairs images with labels",
            "4. **Generate** training data in ShareGPT conversation format",
            "5. **Configure** training parameters for the selected model",
            "",
        ])

        if processing:
            prompt_parts.extend([
                f"## Processing Mode: {processing.task_type.upper()}",
                f"- Matching strategy: {processing.matching_strategy}",
            ])
            if processing.special_instructions:
                prompt_parts.append("- Special notes:")
                for inst in processing.special_instructions:
                    prompt_parts.append(f"  - {inst}")

        if schema:
            prompt_parts.extend([
                "",
                "## Expected Output Schema",
                "```json",
                schema.sample_output[:500],
                "```",
            ])

        prompt_parts.extend([
            "",
            "## Files to Generate",
            "1. `data_loader.py` - PyTorch dataset for image-text pairs",
            "2. `prepare_data.py` - Script to prepare training JSONL",
            "3. `train.py` - Training script with LoRA configuration",
            "4. `config.yaml` - LLaMA-Factory compatible config",
            "",
            "Begin by examining the images and data structure.",
        ])

        return '\n'.join(prompt_parts)

    def _is_schema_file(self, path: Path, content: Any) -> bool:
        """Check if a JSON file is a schema definition."""

        name = path.name.lower()

        # Check filename
        schema_keywords = ['schema', 'output', 'format', 'template', 'spec', 'definition']
        if any(kw in name for kw in schema_keywords):
            return True

        # Check content structure
        if isinstance(content, dict):
            # JSON Schema indicators
            if any(k in content for k in ['$schema', 'type', 'properties', 'required', 'items']):
                return True

        return False

    def _is_archive(self, path: Path) -> bool:
        """Check if path is an archive file."""
        return path.suffix.lower() in self.ARCHIVE_EXTENSIONS or path.name.endswith('.tar.gz')

    def _get_image_dimensions(self, path: Path) -> Optional[Tuple[int, int]]:
        """Get image dimensions."""
        try:
            from PIL import Image
            with Image.open(path) as img:
                return img.size
        except:
            pass
        return None

    async def _extract_archive(self, path: Path) -> Path:
        """Extract archive to temp folder."""

        extract_dir = self.upload_dir / f"extracted_{path.stem}"
        extract_dir.mkdir(exist_ok=True)

        suffix = path.suffix.lower()
        name = path.name.lower()

        if suffix == '.zip':
            with zipfile.ZipFile(path, 'r') as zf:
                zf.extractall(extract_dir)

        elif suffix in ('.tar', '.gz', '.tgz') or name.endswith('.tar.gz'):
            mode = 'r:*'
            if suffix == '.gz' or name.endswith('.tar.gz'):
                mode = 'r:gz'

            with tarfile.open(path, mode) as tf:
                tf.extractall(extract_dir)

        return extract_dir


# Global instance
input_detector = InputDetector()
